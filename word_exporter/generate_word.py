# -*- coding: utf-8 -*-
"""
诚信核查插件工具包 - Word 导出程序

支持的截图结构：
1. 新结构：项目文件夹/核查主体名称/screenshots/*.png
2. 旧结构：项目文件夹/screenshots/*.png

输入项目总文件夹时，程序会为每个核查主体分别生成 Word。
"""

from __future__ import annotations

import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, UnidentifiedImageError


@dataclass
class ExportTarget:
    project_folder: Path
    subject_folder: Path
    screenshots_folder: Path
    report_folder: Path
    project_name: str
    subject_name: str


def main() -> None:
    configure_console()

    print("=" * 60)
    print("诚信核查截图 Word 生成工具")
    print("=" * 60)

    input_path = get_project_folder_from_user()

    if not input_path:
        print("未输入路径，程序已退出。")
        return

    targets = resolve_export_targets(input_path)

    if not targets:
        print("未找到可导出的截图。")
        return

    sites = load_sites_config()
    print(f"\n共识别到 {len(targets)} 个核查主体。")

    generated_files: List[Path] = []

    for target in targets:
        output_file = export_target(target, sites)
        if output_file:
            generated_files.append(output_file)

    if not generated_files:
        print("\n没有生成 Word 文档，请检查截图文件是否有效。")
        return

    print("\nWord 文档已生成：")
    for output_file in generated_files:
        print(output_file)

    print("\n完成。")


def configure_console() -> None:
    if os.name != "nt":
        return

    os.system("chcp 65001 >nul")

    for stream in (sys.stdin, sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            try:
                stream.reconfigure(encoding="utf-8")
            except OSError:
                pass


# ==============================
# 用户输入与目标识别
# ==============================

def get_project_folder_from_user() -> Optional[Path]:
    if len(sys.argv) > 1:
        return Path(" ".join(sys.argv[1:]).strip().strip('"').strip("'"))

    print("\n请输入项目文件夹路径。")
    print("示例：C:\\Users\\你的用户名\\Downloads\\诚信核查插件工具包\\测试项目")
    print("也可以拖入项目文件夹、核查主体文件夹或 screenshots 文件夹，然后按回车。")
    print("也可以把文件夹直接拖到 EXE 图标上运行。")

    raw_path = input("\n路径：").strip()

    if not raw_path:
        return None

    raw_path = raw_path.strip('"').strip("'")
    return Path(raw_path)


def resolve_export_targets(input_path: Path) -> List[ExportTarget]:
    path = input_path.expanduser().resolve()

    if not path.exists():
        print(f"路径不存在：{path}")
        return []

    if path.is_file():
        print("请输入文件夹路径，不要输入单个截图文件。")
        return []

    if path.name.lower() == "screenshots":
        return [build_single_target_from_screenshots(path)]

    child_targets = find_subject_targets_in_project(path)
    if child_targets:
        return child_targets

    direct_screenshots = path / "screenshots"
    if direct_screenshots.exists() and direct_screenshots.is_dir():
        return [build_single_target_from_subject_folder(path)]

    nested_screenshots = [
        item
        for item in path.rglob("screenshots")
        if item.is_dir() and list_screenshots(item)
    ]

    if len(nested_screenshots) == 1:
        print(f"已自动识别截图文件夹：{nested_screenshots[0]}")
        return [build_single_target_from_screenshots(nested_screenshots[0])]

    if len(nested_screenshots) > 1:
        return [build_single_target_from_screenshots(item) for item in sorted(nested_screenshots)]

    print("未找到 screenshots 文件夹。")
    return []


def find_subject_targets_in_project(project_folder: Path) -> List[ExportTarget]:
    targets: List[ExportTarget] = []

    for child in sorted(project_folder.iterdir(), key=lambda item: item.name):
        if not child.is_dir():
            continue

        screenshots_folder = child / "screenshots"
        if screenshots_folder.exists() and screenshots_folder.is_dir() and list_screenshots(screenshots_folder):
            subject_name = child.name
            targets.append(ExportTarget(
                project_folder=project_folder,
                subject_folder=child,
                screenshots_folder=screenshots_folder,
                report_folder=child / "report",
                project_name=project_folder.name,
                subject_name=subject_name,
            ))

    return targets


def build_single_target_from_screenshots(screenshots_folder: Path) -> ExportTarget:
    subject_folder = screenshots_folder.parent
    project_folder = infer_project_folder_from_subject_folder(subject_folder)
    subject_name = infer_subject_name(list_screenshots(screenshots_folder), load_sites_config()) or subject_folder.name

    return ExportTarget(
        project_folder=project_folder,
        subject_folder=subject_folder,
        screenshots_folder=screenshots_folder,
        report_folder=subject_folder / "report",
        project_name=project_folder.name,
        subject_name=subject_name,
    )


def build_single_target_from_subject_folder(subject_folder: Path) -> ExportTarget:
    screenshots_folder = subject_folder / "screenshots"
    project_folder = infer_project_folder_from_subject_folder(subject_folder)
    subject_name = infer_subject_name(list_screenshots(screenshots_folder), load_sites_config()) or subject_folder.name

    return ExportTarget(
        project_folder=project_folder,
        subject_folder=subject_folder,
        screenshots_folder=screenshots_folder,
        report_folder=subject_folder / "report",
        project_name=project_folder.name,
        subject_name=subject_name,
    )


def infer_project_folder_from_subject_folder(subject_folder: Path) -> Path:
    if subject_folder.parent.name in {"诚信核查插件工具包", "诚信核查截图整理工具"}:
        return subject_folder

    return subject_folder.parent


# ==============================
# 导出单个主体
# ==============================

def export_target(target: ExportTarget, sites: List[Dict[str, str]]) -> Optional[Path]:
    screenshots = list_screenshots(target.screenshots_folder)

    if not screenshots:
        print(f"跳过：{target.subject_name}，截图文件夹为空。")
        return None

    valid_screenshots = filter_valid_screenshots(screenshots)

    if not valid_screenshots:
        print(f"跳过：{target.subject_name}，没有有效图片。")
        return None

    subject_name = infer_subject_name(valid_screenshots, sites) or target.subject_name
    entries = build_screenshot_entries(valid_screenshots, sites, subject_name)

    target.report_folder.mkdir(parents=True, exist_ok=True)
    output_file = target.report_folder / f"{sanitize_filename(subject_name)}_诚信核查截图汇总.docx"

    doc = build_document(
        entries=entries,
        project_name=target.project_name,
        subject_name=subject_name,
    )

    doc.save(output_file)
    return output_file


# ==============================
# 读取截图和网站配置
# ==============================

def list_screenshots(screenshots_folder: Path) -> List[Path]:
    if not screenshots_folder.exists():
        return []

    screenshots = [
        path
        for path in screenshots_folder.iterdir()
        if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg"}
    ]

    return sorted(screenshots, key=get_screenshot_sort_key)


def filter_valid_screenshots(screenshots: List[Path]) -> List[Path]:
    valid_screenshots: List[Path] = []

    for screenshot in screenshots:
        try:
            with Image.open(screenshot) as image:
                image.verify()
            valid_screenshots.append(screenshot)
        except (OSError, UnidentifiedImageError) as error:
            print(f"跳过无法读取的图片：{screenshot.name}（{error}）")

    return valid_screenshots


def get_screenshot_sort_key(path: Path) -> tuple:
    match = re.match(r"^(\d+)_", path.name)
    if match:
        return (int(match.group(1)), path.name)

    return (9999, path.name)


def load_sites_config() -> List[Dict[str, str]]:
    sites_file = get_resource_path("extension/sites.json")

    if not sites_file.exists():
        print("未找到网站配置 extension/sites.json，将仅根据截图文件名生成网站标题。")
        return []

    with sites_file.open("r", encoding="utf-8") as f:
        config = json.load(f)

    sites = config.get("sites", [])
    profiles = config.get("profiles", [])
    active_profile_id = config.get("activeProfile", "default")
    active_profile = next(
        (profile for profile in profiles if profile.get("id") == active_profile_id),
        None,
    )

    if not active_profile or not isinstance(active_profile.get("siteIds"), list):
        return sorted(sites, key=lambda item: int(item.get("order", 9999)))

    site_by_id = {
        site.get("id"): site
        for site in sites
        if site.get("id")
    }

    return [
        site_by_id[site_id]
        for site_id in active_profile["siteIds"]
        if site_id in site_by_id and site_by_id[site_id].get("enabled") is not False
    ]


def get_resource_path(relative_path: str) -> Path:
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / relative_path

    return Path(__file__).resolve().parents[1] / relative_path


def build_screenshot_entries(
    screenshots: List[Path],
    sites: List[Dict[str, str]],
    subject_name: str = "",
) -> List[Dict[str, str]]:
    entries: List[Dict[str, str]] = []

    for position, screenshot_path in enumerate(screenshots, start=1):
        file_index = parse_screenshot_index(screenshot_path.name) or position
        filename_site_name = parse_site_name_from_filename(
            screenshot_path.name,
            subject_name,
        )
        site = find_site_by_name(filename_site_name, sites)

        entries.append({
            "siteIndex": str(file_index),
            # 截图时使用的网站名称已经固化在文件名中。不能再按序号读取
            # sites.json，否则自定义清单发生增删或重排后会把标题配给错误截图。
            "siteName": site.get("name") or filename_site_name,
            "url": site.get("url", ""),
            "capturedAt": parse_captured_at_from_filename(screenshot_path.name),
            "screenshotPath": str(screenshot_path),
        })

    return entries


def infer_subject_name(screenshots: List[Path], sites: List[Dict[str, str]]) -> str:
    for screenshot_path in screenshots:
        subject_name = infer_subject_name_from_filename(screenshot_path.name, sites)
        if subject_name:
            return subject_name

    return ""


def infer_subject_name_from_filename(filename: str, sites: List[Dict[str, str]]) -> str:
    file_index = parse_screenshot_index(filename)
    if not file_index or not (0 < file_index <= len(sites)):
        return ""

    stem = Path(filename).stem
    stem = re.sub(r"^\d+_", "", stem)
    stem = re.sub(r"_\d{8}_\d{6}$", "", stem)

    safe_site_name = sanitize_file_part(sites[file_index - 1].get("name", ""))
    prefix = f"{safe_site_name}_"

    if not safe_site_name or not stem.startswith(prefix):
        return ""

    return stem[len(prefix):].replace("_", " ").strip()


def parse_screenshot_index(filename: str) -> Optional[int]:
    match = re.match(r"^(\d+)_", filename)
    if not match:
        return None

    return int(match.group(1))


def parse_site_name_from_filename(filename: str, subject_name: str = "") -> str:
    stem = Path(filename).stem
    stem = re.sub(r"^\d+_", "", stem)
    stem = re.sub(r"_\d{8}_\d{6}$", "", stem)

    safe_subject_name = sanitize_file_part(subject_name)
    subject_suffix = f"_{safe_subject_name}"
    if safe_subject_name and stem.endswith(subject_suffix):
        stem = stem[:-len(subject_suffix)]

    return stem.replace("_", " ") or "未命名网站"


def find_site_by_name(
    site_name: str,
    sites: List[Dict[str, str]],
) -> Dict[str, str]:
    """按截图中固化的网站名匹配配置，仅用于补充原名称和网址。"""
    normalized_name = sanitize_file_part(site_name).casefold()
    if not normalized_name:
        return {}

    for site in sites:
        configured_name = str(site.get("name", ""))
        if sanitize_file_part(configured_name).casefold() == normalized_name:
            return site

    return {}


def parse_captured_at_from_filename(filename: str) -> str:
    match = re.search(r"_(\d{8})_(\d{6})\.[^.]+$", filename)
    if not match:
        return ""

    date_text, time_text = match.groups()
    return (
        f"{date_text[0:4]}-{date_text[4:6]}-{date_text[6:8]} "
        f"{time_text[0:2]}:{time_text[2:4]}:{time_text[4:6]}"
    )


# ==============================
# 生成 Word
# ==============================

def build_document(
    entries: List[Dict[str, str]],
    project_name: str,
    subject_name: str,
) -> Document:
    doc = Document()

    setup_document_layout(doc)
    setup_styles(doc)

    add_cover_title(doc, project_name, subject_name)
    add_summary_table(doc, entries, subject_name)
    add_screenshot_sections(doc, entries, subject_name)

    return doc


def setup_document_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT

    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    section.page_width = Cm(21)
    section.page_height = Cm(29.7)


def setup_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_cover_title(doc: Document, project_name: str, subject_name: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run("诚信核查截图汇总")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(20)

    doc.add_paragraph()

    info_items = [
        ("项目名称", project_name),
        ("核查主体", subject_name),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = True

    for label, value in info_items:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

        set_cell_shading(row.cells[0], "F2F2F2")
        set_cell_font(row.cells[0], bold=True)
        set_cell_font(row.cells[1], bold=False)

    doc.add_paragraph()


def add_summary_table(doc: Document, entries: List[Dict[str, str]], subject_name: str) -> None:
    heading = doc.add_paragraph()
    heading_run = heading.add_run("一、截图清单")
    heading_run.bold = True
    heading_run.font.name = "黑体"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading_run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["序号", "核查网站", "核查主体", "截图时间"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        set_cell_font(table.rows[0].cells[i], bold=True)

    for item in entries:
        row = table.add_row()
        row.cells[0].text = item.get("siteIndex", "")
        row.cells[1].text = item.get("siteName", "")
        row.cells[2].text = subject_name
        row.cells[3].text = item.get("capturedAt", "")

        for cell in row.cells:
            set_cell_font(cell, bold=False)

    doc.add_paragraph()


def add_screenshot_sections(
    doc: Document,
    entries: List[Dict[str, str]],
    subject_name: str,
) -> None:
    heading = doc.add_paragraph()
    heading_run = heading.add_run("二、核查截图")
    heading_run.bold = True
    heading_run.font.name = "黑体"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading_run.font.size = Pt(14)

    for idx, item in enumerate(entries):
        site_index = item.get("siteIndex", str(idx + 1))
        site_name = item.get("siteName", "未命名网站")
        captured_at = item.get("capturedAt", "")
        url = item.get("url", "")
        screenshot_path = Path(item.get("screenshotPath", ""))

        add_site_block(
            doc=doc,
            site_index=site_index,
            site_name=site_name,
            subject_name=subject_name,
            captured_at=captured_at,
            url=url,
            screenshot_path=screenshot_path,
        )

        if idx != len(entries) - 1:
            doc.add_page_break()


def add_site_block(
    doc: Document,
    site_index: str,
    site_name: str,
    subject_name: str,
    captured_at: str,
    url: str,
    screenshot_path: Optional[Path],
) -> None:
    title = doc.add_paragraph()
    title_run = title.add_run(f"{site_index}. {site_name}")
    title_run.bold = True
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_run.font.size = Pt(13)

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"

    meta_items = [
        ("核查网站", site_name),
        ("核查主体", subject_name),
        ("核查时间", captured_at),
        ("页面链接", url),
    ]

    for label, value in meta_items:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

        set_cell_shading(row.cells[0], "F2F2F2")
        set_cell_font(row.cells[0], bold=True)
        set_cell_font(row.cells[1], bold=False)

    doc.add_paragraph()

    if screenshot_path and screenshot_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(str(screenshot_path), width=get_available_image_width(doc))
    else:
        warning = doc.add_paragraph()
        warning_run = warning.add_run("【未找到对应截图文件】")
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(192, 0, 0)


# ==============================
# Word 样式辅助函数
# ==============================

def set_cell_font(cell, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10)
            run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def get_available_image_width(doc: Document) -> Cm:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


# ==============================
# 文件名清理
# ==============================

def sanitize_filename(name: str) -> str:
    return sanitize_file_part(name) or "诚信核查截图汇总"


def sanitize_file_part(name: str) -> str:
    invalid_chars = '\\/:*?"<>|'
    result = "".join("_" if ch in invalid_chars else ch for ch in name)
    result = re.sub(r"\s+", "_", result.strip())
    result = re.sub(r"_+", "_", result)
    return result[:80]


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n用户已取消。")
        sys.exit(1)
    except Exception as e:
        print(f"\n程序运行出错：{e}")
        sys.exit(1)
